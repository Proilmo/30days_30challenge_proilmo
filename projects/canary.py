#!/usr/bin/env python3
"""
canary.py — a complete, self-contained canary token system.

Everything from Phase 1 and Phase 2 lives here in one file, organized
into clearly labeled sections so it's still easy to navigate:

    1. DATABASE        - SQLite setup, tokens + events tables
    2. TOKEN GENERATOR  - create/list/fetch tokens
    3. ALERTS           - email / webhook notifications on trigger
    4. PACKAGER         - QR code / Word doc / HTML bait file generation
    5. LISTENER SERVER  - the Flask app that catches triggers
    6. CLI              - command-line entry point tying it all together

Usage:
    python canary.py init
    python canary.py create --name "finance.docx" --type doc --package docx
    python canary.py serve                     # run the listener (separate terminal)
    python canary.py list
    python canary.py events

Dependencies:
    pip install flask requests qrcode python-docx pillow
"""

import argparse
import os
import smtplib
import sqlite3
import uuid
from contextlib import contextmanager
from datetime import datetime, timezone
from email.mime.text import MIMEText


# ===========================================================================
# 1. DATABASE
# ===========================================================================

DB_PATH = "canary.db"


def init_db():
    """Create the tokens and events tables if they don't already exist."""
    with get_connection() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS tokens (
                id          TEXT PRIMARY KEY,
                name        TEXT NOT NULL,
                token_type  TEXT NOT NULL,
                created_at  TEXT NOT NULL,
                notes       TEXT
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS events (
                id           INTEGER PRIMARY KEY AUTOINCREMENT,
                token_id     TEXT NOT NULL,
                triggered_at TEXT NOT NULL,
                source_ip    TEXT,
                user_agent   TEXT,
                referrer     TEXT,
                FOREIGN KEY (token_id) REFERENCES tokens (id)
            )
        """)
    print("[db] tables ready")


@contextmanager
def get_connection():
    """Open a SQLite connection, auto-commit on success, always close."""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        yield conn
        conn.commit()
    finally:
        conn.close()


# ===========================================================================
# 2. TOKEN GENERATOR
# ===========================================================================

# Change this to your real public domain when deploying for real, e.g.
# "https://canary.yourdomain.com"
BASE_URL = "http://127.0.0.1:5000"


def create_token(name: str, token_type: str = "url", notes: str = "") -> dict:
    """Create a new token, store it, and return its trigger_url."""
    token_id = str(uuid.uuid4())
    created_at = datetime.now(timezone.utc).isoformat()

    with get_connection() as conn:
        conn.execute(
            "INSERT INTO tokens (id, name, token_type, created_at, notes) "
            "VALUES (?, ?, ?, ?, ?)",
            (token_id, name, token_type, created_at, notes),
        )

    return {
        "id": token_id,
        "name": name,
        "type": token_type,
        "created_at": created_at,
        "trigger_url": f"{BASE_URL}/t/{token_id}",
    }


def list_tokens() -> list[dict]:
    with get_connection() as conn:
        rows = conn.execute("SELECT * FROM tokens ORDER BY created_at DESC").fetchall()
    return [dict(row) for row in rows]


def get_token(token_id: str) -> dict | None:
    with get_connection() as conn:
        row = conn.execute("SELECT * FROM tokens WHERE id = ?", (token_id,)).fetchone()
    return dict(row) if row else None


# ===========================================================================
# 3. ALERTS
# ===========================================================================
#
# SMTP config (all optional -- if unset, alerts just print to console):
#   export CANARY_SMTP_HOST="smtp.gmail.com"
#   export CANARY_SMTP_PORT="587"
#   export CANARY_SMTP_USER="you@gmail.com"
#   export CANARY_SMTP_PASS="your-app-password"
#   export CANARY_ALERT_TO="you@gmail.com"
#
# Webhook config (optional, for Slack/Discord instead of email):
#   export CANARY_WEBHOOK_URL="https://hooks.slack.com/services/..."

def send_email_alert(token_name, token_id, source_ip, user_agent, triggered_at):
    host = os.environ.get("CANARY_SMTP_HOST")
    port = os.environ.get("CANARY_SMTP_PORT")
    user = os.environ.get("CANARY_SMTP_USER")
    password = os.environ.get("CANARY_SMTP_PASS")
    to_addr = os.environ.get("CANARY_ALERT_TO")

    if not all([host, port, user, password, to_addr]):
        print(f"[ALERT] Token '{token_name}' ({token_id}) triggered at "
              f"{triggered_at} from {source_ip} ({user_agent})")
        return

    body = (
        f"Your canary token was triggered.\n\n"
        f"Token name: {token_name}\nToken ID:   {token_id}\n"
        f"Time:       {triggered_at}\nSource IP:  {source_ip}\n"
        f"User-Agent: {user_agent}\n"
    )
    msg = MIMEText(body)
    msg["Subject"] = f"Canary triggered: {token_name}"
    msg["From"] = user
    msg["To"] = to_addr

    try:
        with smtplib.SMTP(host, int(port)) as server:
            server.starttls()
            server.login(user, password)
            server.send_message(msg)
        print(f"[alerts] email sent for token {token_id}")
    except Exception as e:
        print(f"[alerts] FAILED to send email: {e}")


def send_webhook_alert(token_name, token_id, source_ip, user_agent, triggered_at):
    webhook_url = os.environ.get("CANARY_WEBHOOK_URL")
    if not webhook_url:
        return
    import requests
    text = (
        f"Canary triggered: {token_name}\nToken ID: {token_id}\n"
        f"Time: {triggered_at}\nSource IP: {source_ip}\nUser-Agent: {user_agent}"
    )
    try:
        requests.post(webhook_url, json={"text": text}, timeout=5)
        print(f"[alerts] webhook sent for token {token_id}")
    except Exception as e:
        print(f"[alerts] FAILED to send webhook: {e}")


def notify(token_name, token_id, source_ip, user_agent, triggered_at):
    send_email_alert(token_name, token_id, source_ip, user_agent, triggered_at)
    send_webhook_alert(token_name, token_id, source_ip, user_agent, triggered_at)


# ===========================================================================
# 4. PACKAGER — QR code / Word doc / HTML bait file generation
# ===========================================================================

def generate_qr(trigger_url: str, output_path: str = "token.png"):
    """Encode trigger_url as a scannable QR code PNG."""
    import qrcode
    img = qrcode.make(trigger_url)
    img.save(output_path)
    print(f"[packager] QR token saved to {output_path}")
    return output_path


def generate_docx(trigger_url: str, output_path: str = "token.docx",
                   title: str = "Confidential"):
    """
    Build a .docx that fetches trigger_url when opened in Word, using an
    INCLUDEPICTURE field code (python-docx's add_picture() only embeds
    local bytes, so we build the field manually via raw OOXML).
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(
        "This document contains sensitive information. "
        "Do not distribute without authorization."
    )

    p = doc.add_paragraph()

    # --- fldChar begin ---
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    # --- instrText: the actual field instruction ---
    run = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' INCLUDEPICTURE "{trigger_url}" \\* MERGEFORMAT '
    run._r.append(instr)

    # --- fldChar separate ---
    run = p.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    # --- placeholder shown until Word evaluates the field ---
    p.add_run("[loading attachment preview...]")

    # --- fldChar end ---
    run = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    doc.add_paragraph(
        "\nIf the image above does not display, enable content and "
        "external links when prompted by Word."
    )

    # Nudge Word to evaluate fields on open (not guaranteed -- modern
    # Word may still prompt or block external content by default).
    settings_el = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings_el.insert(0, update_fields)

    doc.save(output_path)
    print(f"[packager] docx token saved to {output_path}")
    return output_path


def generate_html(trigger_url: str, output_path: str = "token.html",
                   subject: str = "Quarterly Report"):
    """Build an HTML file with the token as an invisible <img> web bug."""
    html = f"""<!DOCTYPE html>
<html>
<head><title>{subject}</title></head>
<body>
<p>This is an automated notice regarding: {subject}</p>
<img src="{trigger_url}" width="1" height="1" style="display:block;border:0;" alt="">
</body>
</html>
"""
    with open(output_path, "w") as f:
        f.write(html)
    print(f"[packager] HTML token saved to {output_path}")
    return output_path


def package_token(token: dict, kind: str, out_prefix: str = None):
    """
    Convenience wrapper: given a token dict (from create_token) and a
    kind ('qr' | 'docx' | 'html' | 'all'), generate the requested bait
    file(s) named after the token's id.
    """
    prefix = out_prefix or token["id"][:8]
    url = token["trigger_url"]
    paths = []
    if kind in ("qr", "all"):
        paths.append(generate_qr(url, f"{prefix}.png"))
    if kind in ("docx", "all"):
        paths.append(generate_docx(url, f"{prefix}.docx"))
    if kind in ("html", "all"):
        paths.append(generate_html(url, f"{prefix}.html"))
    return paths


# ===========================================================================
# 5. LISTENER SERVER — catches triggers
# ===========================================================================

# 1x1 transparent GIF, served on every trigger regardless of outcome.
TRANSPARENT_PIXEL = (
    b"\x47\x49\x46\x38\x39\x61\x01\x00\x01\x00\x80\x00\x00\xff\xff\xff"
    b"\x00\x00\x00\x21\xf9\x04\x01\x00\x00\x00\x00\x2c\x00\x00\x00\x00"
    b"\x01\x00\x01\x00\x00\x02\x02\x44\x01\x00\x3b"
)


def create_app():
    """
    Build and return the Flask app. Wrapped in a function (rather than
    a bare module-level `app = Flask(__name__)`) so that importing this
    file for its other functions (create_token, generate_qr, etc.)
    doesn't require Flask to even be installed unless you actually run
    `python canary.py serve`.
    """
    from flask import Flask, request, Response

    app = Flask(__name__)

    @app.route("/t/<token_id>", methods=["GET"])
    def trigger(token_id):
        token = get_token(token_id)
        if token:
            _log_and_alert(token, request)
        # Same response whether the token is real or not, and always fast --
        # never tip off an attacker that they've hit a trap.
        return Response(TRANSPARENT_PIXEL, mimetype="image/gif")

    return app


def _log_and_alert(token: dict, req):
    triggered_at = datetime.now(timezone.utc).isoformat()
    source_ip = _request_ip(req)
    user_agent = req.headers.get("User-Agent")
    referrer = req.headers.get("Referer")

    with get_connection() as conn:
        conn.execute(
            "INSERT INTO events (token_id, triggered_at, source_ip, user_agent, referrer) "
            "VALUES (?, ?, ?, ?, ?)",
            (token["id"], triggered_at, source_ip, user_agent, referrer),
        )

    notify(
        token_name=token["name"],
        token_id=token["id"],
        source_ip=source_ip,
        user_agent=user_agent,
        triggered_at=triggered_at,
    )


def _request_ip(req) -> str:
    """Prefer X-Forwarded-For (set by a trusted reverse proxy) over remote_addr."""
    forwarded = req.headers.get("X-Forwarded-For")
    if forwarded:
        return forwarded.split(",")[0].strip()
    return req.remote_addr


# ===========================================================================
# 6. CLI
# ===========================================================================

def cmd_init(args):
    init_db()


def cmd_create(args):
    token = create_token(name=args.name, token_type=args.type, notes=args.notes or "")
    print("Token created:")
    print(f"  ID:          {token['id']}")
    print(f"  Name:        {token['name']}")
    print(f"  Type:        {token['type']}")
    print(f"  Trigger URL: {token['trigger_url']}")

    if args.package:
        print()
        paths = package_token(token, kind=args.package)
        print("Bait files:")
        for p in paths:
            print(f"  {p}")


def cmd_list(args):
    tokens = list_tokens()
    if not tokens:
        print("No tokens created yet. Use 'python canary.py create ...' first.")
        return
    for t in tokens:
        print(f"[{t['token_type']:6}] {t['id']}  {t['name']}  (created {t['created_at']})")


def cmd_events(args):
    with get_connection() as conn:
        rows = conn.execute("""
            SELECT events.triggered_at, events.source_ip, events.user_agent,
                   tokens.name AS token_name
            FROM events JOIN tokens ON tokens.id = events.token_id
            ORDER BY events.triggered_at DESC
        """).fetchall()
    if not rows:
        print("No trigger events yet.")
        return
    for r in rows:
        print(f"{r['triggered_at']}  |  {r['token_name']:25}  |  "
              f"{r['source_ip']:15}  |  {r['user_agent']}")


def cmd_serve(args):
    init_db()
    app = create_app()
    print(f"[app] canary listener running on {BASE_URL}")
    app.run(host="0.0.0.0", port=5000, debug=False)


def main():
    parser = argparse.ArgumentParser(description="Canary token manager (single-file edition)")
    subparsers = parser.add_subparsers(dest="command", required=True)

    subparsers.add_parser("init", help="Set up the database").set_defaults(func=cmd_init)

    p_create = subparsers.add_parser("create", help="Create a new token")
    p_create.add_argument("--name", required=True, help="Human-readable label for this token")
    p_create.add_argument("--type", default="url", help="Token type label, e.g. url/doc/qrcode")
    p_create.add_argument("--notes", default="", help="Optional notes")
    p_create.add_argument(
        "--package", choices=["qr", "docx", "html", "all"], default=None,
        help="Also generate a bait file of this kind"
    )
    p_create.set_defaults(func=cmd_create)

    subparsers.add_parser("list", help="List all tokens").set_defaults(func=cmd_list)
    subparsers.add_parser("events", help="List all trigger events").set_defaults(func=cmd_events)
    subparsers.add_parser("serve", help="Run the listener server").set_defaults(func=cmd_serve)

    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
